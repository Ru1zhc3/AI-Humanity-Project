from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\CourseProject\cs4501_26spring_final_project")
OUT = ROOT / "Political Content Moderation Across Different Models and Languages_script.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.name = "Aptos"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "E6F2EA")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False)
    doc.add_paragraph()


def add_section(doc: Document, title: str, script: list[str], table: tuple[list[str], list[list[str]]] | None = None) -> None:
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(24, 50, 38)
    if table:
        add_table(doc, table[0], table[1])
    for para in script:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(para)
        run.font.name = "Aptos"
        run.font.size = Pt(11)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Political Content Moderation Across Different Models and Languages")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Aptos Display"
    title_run.font.color.rgb = RGBColor(24, 50, 38)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Speaker Script for Results, Outcomes, and Conclusion")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.name = "Aptos"
    subtitle_run.font.color.rgb = RGBColor(78, 103, 88)
    doc.add_paragraph()

    add_section(
        doc,
        "Result Slide 1 - Translation Effects (P16 in current slides)",
        [
            "On this slide, we are looking only at the original, non-perturbed statements. For each non-English version, I compare the model's political-lean label against that same model's English label for the same source statement. So a shifted lean means the same underlying statement was recoded into a different political-compass category after translation.",
            "The first important result is that Gemma shifts more than Codex in every non-English language. Codex has an average translation shift of 19.5%, while Gemma has an average translation shift of 34.8%. In other words, Gemma is much more sensitive to translation even when the underlying statement is intended to stay semantically equivalent.",
            "For Codex, the largest translation shift is Farsi at 23.6%, and the smallest is Arabic at 16.0%. For Gemma, the largest translation shift is Amharic at 42.2%, and the smallest non-English shift is 32.0%, which appears for Russian, Farsi, and Latin American Spanish.",
            "The directionality table adds a second layer: not only whether labels changed, but where the changed rows moved. Using our shifted-row standard, the two languages with exact same-direction behavior across both models are French, where shifted rows most often move to Lib-Right, and Russian, where shifted rows most often move to Lib-Left. Most other languages are model-dependent, which means the translation effect is structured but not universal.",
            "A possible interpretation is that translation changes the political vocabulary, cultural context, or connotation available to the model. Another possibility is uneven multilingual training coverage. The key point for the presentation is that translation is not just a preprocessing step; it is part of the evaluation condition.",
        ],
        (
            ["Model", "Average", "Maximum", "Minimum excluding English"],
            [
                ["GPT-5.4-mini", "19.5%", "Farsi, 23.6%", "Arabic, 16.0%"],
                ["Gemma 4 e4b", "34.8%", "Amharic, 42.2%", "Russian/Farsi/LatAm Spanish, 32.0%"],
            ],
        ),
    )

    add_section(
        doc,
        "Result Slide 2 - Perturbation Effects (P17 in current slides)",
        [
            "This slide looks at perturbations, which means different framing versions of the same statement. Here, shifted lean is measured against the majority lean across the 14 perturbation variants for the same model, language, and source statement. This lets us ask whether changing the prompt form changes the political stance the model infers.",
            "For GPT-5.4-mini, the average perturbation shift is 23.9%. Its maximum is Open Ended at 40.1%, and its minimum is Active Voice at 15.6%. For Gemma, the average perturbation shift is higher, 30.7%. Its maximum is also Open Ended at 47.3%, and its minimum is also Active Voice at 17.3%. So both models agree on the most and least destabilizing perturbation type, but Gemma moves more overall.",
            "Model agreement means the percentage of rows where Codex and Gemma assign the exact same political-lean label for the same statement, language, and perturbation. Under perturbation, that agreement is 59.5%, which is lower than the 64.6% agreement on the non-perturbed translation rows. So framing makes the two models diverge more.",
            "Using the dominant shifted-row standard, the perturbations with exact same-direction behavior across both models are Open Ended, Yes-No Forced Choice, Formal, and Hedged, which all move most often to Centrist, plus Loss Framing, which moves most often to Lib-Left. This is useful because it shows that some framing types have shared directional pressure across model families.",
            "For GPT-5.4-mini specifically, the overall perturbation pattern is left or centrist. Across all 14 perturbation types, the dominant shifted lean is Auth-Left for six types, Lib-Left for three types, and Centrist for five types; none of the perturbation types has a dominant right-shift. That is why I would describe Codex's perturbation behavior as left/centrist rather than neutral.",
        ],
        (
            ["Model", "Average", "Maximum", "Minimum"],
            [
                ["GPT-5.4-mini", "23.9%", "Open Ended, 40.1%", "Active Voice, 15.6%"],
                ["Gemma 4 e4b", "30.7%", "Open Ended, 47.3%", "Active Voice, 17.3%"],
            ],
        ),
    )

    add_section(
        doc,
        "Outcome Slide",
        [
            "The outcome is not simply that labels sometimes change. The deeper outcome is that the changes have structure. Translation, perturbation, and model family each act like independent evaluation variables.",
            "First, translation matters. The same base statement can receive different political-lean labels after translation. Gemma is more affected than Codex in every non-English language, which suggests that multilingual evaluation should not assume English is a clean baseline for all languages.",
            "Second, perturbation matters. Open-ended, hedged, formal, and yes-no framing tend to pull shifted rows toward Centrist, while loss framing pulls both models toward Lib-Left. Codex's perturbation behavior is especially important because its dominant movements are left or centrist and never right.",
            "Third, model family matters. Gemma is more sensitive overall, with higher average shift rates for both translation and perturbation. Codex is more stable, but it still shows systematic recoding. That means a single model should not be treated as a neutral judge for political moderation.",
            "One possible reason for the model difference is that the models have different training distributions, alignment objectives, and multilingual coverage. Gemma may be more reactive to surface wording and translation variation, while Codex appears more conservative in label movement but still follows a moderation-like left/centrist pattern under perturbation.",
        ],
    )

    add_section(
        doc,
        "Conclusion Slide",
        [
            "The conclusion has two parts. The first part is about translation and perturbation. Our results show that these are not neutral transformations. Translation can change political vocabulary and cultural cues, and perturbation can change the pragmatic frame of the same statement. Both can move the model's political-lean judgment.",
            "The second part is about model comparison. Gemma and Codex are not interchangeable judges. Gemma shifts more under both translation and perturbation, while Codex is more stable but shows a left/centrist tendency under perturbation. So model choice is itself part of the moderation behavior.",
            "The final takeaway is that multilingual political moderation should be audited across languages, perturbation styles, and model families before claiming neutrality or fairness. A single language, a single prompt form, or a single model cannot represent the whole evaluation space.",
            "This supports our project proposal because it shows implicit alignment or moderation behavior through variation. The goal is not to say one political label is the objectively correct answer; the goal is to show that the same content can be treated differently depending on language, framing, and model.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
